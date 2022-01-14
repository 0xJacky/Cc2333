import os

from docx import *


def get_docx_list(path='./Python'):
    r = []
    for root, dirs, files in os.walk(path):
        for name in files:
            if 'MACOS' in root:
                continue
            if '~$' not in name and 'doc' in name:
                r.append((root, name))
    return r


def read_docx(path):
    is_report = True
    course_name = ''
    exp_name = ''
    teacher_name = ''
    reporter = []
    exp_time = ''
    submit_time = ''

    print(path)
    doc = Document(path)
    try:
        tmp = ''
        for _p in doc.paragraphs[0:min(10, len(doc.paragraphs))]:
            tmp += _p.text
        tmp = tmp.replace(' ', '')
        print(tmp)
        if '实验报告' not in tmp:
            is_report = False
    except:
        is_report = False

    if is_report:
        _para = doc.tables[0].cell(0, 0).paragraphs
        print('总共有 %s 段文字' % (len(_para)))
        line_cnt = 0
        for i in _para:
            print(line_cnt, i.text)
            line_cnt += 1
        # 课程名称
        course_name = _para[0].text[5:].strip()
        print('课程名', course_name)
        # 实验名称
        exp_name = _para[1].text[5:].strip()
        print('实验名', exp_name)
        # 指导老师
        teacher_name = _para[4].text[5:].strip()
        print('老师', teacher_name)
        # 报告人
        reporter = _para[5].text[5:].strip().replace(' ', '').split('学号')
        print(reporter)
        if len(reporter) < 2:
            reporter.append(' ')
        # 实验时间
        exp_time = _para[6].text[5:].strip()
        # exp_time = re.findall(r'[\s]*(\d{4})[\s]*年[\s]*(\d{1,2})[\s]*月[\s]*(\d{1,2})[\s]*日', exp_time)[0]
        print('实验时间', exp_time)
        # 提交时间
        submit_time = _para[7].text[5:].strip()

        print('提交', submit_time)
    return course_name, exp_name, teacher_name, reporter, exp_time, submit_time, is_report


if __name__ == '__main__':
    docx_list = get_docx_list()
    print(docx_list)
    for i in docx_list:
        print(i)
        read_docx(i)
