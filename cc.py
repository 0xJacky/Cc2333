import datetime
from docx import *
from docx.shared import Pt
from docx.enum.text import WD_UNDERLINE
from docx2pdf import convert


def magic(para, line_number, m_text, end_i=5, add_space_len=0):
    print(para[line_number].text[0:end_i])
    max_len = 38
    # 手动调整偏移
    max_len += (add_space_len * 2)
    print(max_len)
    # 居中测试
    content_len = len(m_text)
    space_len = (max_len - end_i - content_len) // 2
    space_len -= 1
    print('content_len, space_len', content_len, space_len)
    # 保留部分
    label_text = para[line_number].text[0:end_i]
    para[line_number].clear()
    run = para[line_number].add_run()
    run.font.bold = True
    run.font.size = Pt(14)
    run.text = label_text

    run = para[line_number].add_run()
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.underline = WD_UNDERLINE.THICK
    if content_len < 15:
        run.text += (' ' * space_len)
    else:
        run.text += ' '
    run.text += m_text
    run.text += (' ' * max_len)


def magic_reporter(para, line_number, name, school_id):
    para[line_number].clear()

    run = para[line_number].add_run()
    run.text = '报 告 人：'
    run.font.bold = True
    run.font.size = Pt(14)

    run = para[line_number].add_run()
    run.text = '    ' + name + '    '
    run.font.underline = WD_UNDERLINE.THICK
    run.font.bold = True
    run.font.size = Pt(14)

    run = para[line_number].add_run()
    run.text = '学号：'
    run.font.bold = True
    run.font.size = Pt(14)

    run = para[line_number].add_run()
    run.text = '  ' + school_id + (' ' * 10)
    run.font.underline = WD_UNDERLINE.THICK
    run.font.bold = True
    run.font.size = Pt(14)


def magic_exp_date(para, date_tuple):
    trans_week = ['一', '二', '三', '四', '五', '六', '日']
    runs = para[7].runs
    idx = 0
    for i in runs:
        print(idx, i.text)
        idx += 1

    date_tuple = tuple(map(int, date_tuple))

    runs[1].text = '  %d  ' % date_tuple[0]
    runs[3].text = '  %d  ' % date_tuple[1]
    runs[5].text = '  %d  ' % date_tuple[2]

    runs[9].text = ' %s ' % trans_week[datetime.date(*date_tuple).today().weekday()]


if __name__ == '__main__':
    cover_path = '1 实验报告封皮.docx'

    cover = Document(cover_path)

    _para = cover.paragraphs
    _para[0].clear()
    run = _para[0].add_run()
    run.text = '课程编号：'
    run.font.size = Pt(14)
    run = _para[0].add_run()
    run.text = '  IB01017  '
    run.font.underline = True

    _para = cover.tables[1].cell(0, 0).paragraphs
    print('总共有 %s 段文字' % (len(_para)))
    line_cnt = 0
    for i in _para:
        print(line_cnt, i.text, i.style.font.size.pt)
        line_cnt += 1

    magic(_para, 0, 'Python 程序设计')
    magic(_para, 1, '初识 Python')
    magic(_para, 2, '2019级物联网五班', end_i=7)
    magic(_para, 3, '柯笑')
    magic(_para, 6, 'C5-428')
    magic_exp_date(_para, ('2021', '12', '28'))
    # magic(_para, 7, '2021年12月23日', add_space_len=-1)

    magic_reporter(_para, 4, '余圳曦', '201904020209')

    cover.save('save.docx')
    convert('save.docx', 'save.pdf')
