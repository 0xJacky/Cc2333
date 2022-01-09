import os
from pathlib import Path
import argparse

from PyPDF2 import PdfFileReader, PdfFileWriter
from docx2pdf import convert

from docx_func import get_docx_list, read_docx
from docx import Document
from docx.shared import Pt
from cc import magic, magic_reporter, magic_exp_date


def change_cover(_id, class_name, location, orig_docx_path, output_dir='./output'):
    # 合成新报告封面
    cover_path = '1 实验报告封皮.docx'
    cover = Document(cover_path)

    _para = cover.paragraphs
    _para[0].clear()
    run = _para[0].add_run()
    run.text = '课程编号：'
    run.font.size = Pt(14)
    run = _para[0].add_run()
    run.text = '  %s  ' % _id
    run.font.underline = True

    _para = cover.tables[1].cell(0, 0).paragraphs
    print('总共有 %s 段文字' % (len(_para)))

    # 提取原报告封面的信息
    course_name, exp_name, teacher_name, reporter, exp_time, submit_time = read_docx(orig_docx_path)
    print(course_name, exp_name, teacher_name, exp_time, submit_time)

    # 课程名称
    magic(_para, 0, course_name)
    # 实验名称
    magic(_para, 1, exp_name)
    # 班级
    magic(_para, 2, class_name, end_i=7)
    # 指导老师
    magic(_para, 3, teacher_name)
    # 报告人
    magic_reporter(_para, 4, reporter[0], reporter[1])
    # 实验地点
    magic(_para, 6, location)
    # 实验时间
    magic_exp_date(_para, exp_time)
    # 提交时间
    magic(_para, 8, submit_time)

    # 设置格式
    _para[0].style.font.size = Pt(10.5)
    _para[0].style.font.bold = None

    # 定义导出路径
    orig_docx_name = Path(orig_docx_path).stem
    docx_path = '%s/%s.docx' % (output_dir, orig_docx_name)
    cover_pdf_path = '%s/%s-cover.pdf' % (output_dir, orig_docx_name)
    orig_pdf_path = '%s/%s-orig.pdf' % (output_dir, orig_docx_name)

    # 保存新封面 docx
    cover.save(docx_path)
    # 转换新封面为 pdf
    convert(docx_path, cover_pdf_path)
    # 转换旧报告为 pdf
    convert(orig_docx_path, '%s/%s-orig.pdf' % (output_dir, orig_docx_name))

    # 选择合并 pdf
    cover_pdf = PdfFileReader(cover_pdf_path)
    orig_pdf = PdfFileReader(orig_pdf_path)
    output = PdfFileWriter()
    p = cover_pdf.getPage(0)
    output.addPage(p)

    print(orig_pdf.getNumPages())

    for i in range(1, orig_pdf.getNumPages()):
        p = orig_pdf.getPage(i)
        output.addPage(p)
    output_pdf_path = '%s/%s.pdf' % (output_dir, orig_docx_name)
    with open(output_pdf_path, 'wb') as f:
        output.write(f)

    # 清理
    os.remove(docx_path)
    os.remove(cover_pdf_path)
    os.remove(orig_pdf_path)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Project Cc2333')
    parser.add_argument('-i', '--id', action="store", help='course id')
    parser.add_argument('-c', '--class_name', action="store", help='class name')
    parser.add_argument('-l', '--location', action="store", help='exp location')
    parser.add_argument('-s', '--source', action="store", help='reports dir')
    parser.add_argument('-o', '--output', action="store", help='output dir')
    args = parser.parse_args()

    docx_list = get_docx_list(args.source)
    print(docx_list)

    for d in docx_list:
        change_cover(args.id, args.class_name, args.location, d, output_dir=args.output)
